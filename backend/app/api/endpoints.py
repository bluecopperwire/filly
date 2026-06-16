"""
FILLY API Endpoints
===================
All routes that the frontend consumes. Mounted under the /api prefix in main.py.

Endpoints:
    POST   /analyze              – Analyze text for normalizations & grammar errors
    POST   /document             – Create a new document
    GET    /document/{id}        – Retrieve a document by ID
    PUT    /document/{id}        – Update a document
    GET    /documents            – List all documents
    POST   /upload               – Upload a .txt or .docx file and extract text
    POST   /suggestion/{id}/accept – Mark a suggestion as accepted
    POST   /suggestion/{id}/ignore – Mark a suggestion as ignored
    GET    /analytics            – Get global analytics
    GET    /analytics/{doc_id}   – Get analytics for a specific document
    GET    /health               – Health check
"""

import logging
from io import BytesIO

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session

from app.database.database import get_db
from app.models.models import Document, Suggestion
from app.schemas.schemas import (
    AnalyzeRequest,
    AnalyzeResponse,
    DocumentCreate,
    DocumentUpdate,
    DocumentResponse,
    UploadResponse,
    SuggestionActionResponse,
    AnalyticsResponse,
    NormalizationStats,
    GrammarStats,
)
from app.services.normalizer import get_normalizer
from app.services.gec import get_gec_service

logger = logging.getLogger(__name__)

router = APIRouter()


# ─── Health ──────────────────────────────────────────────────────────


@router.get("/health")
async def health_check():
    """Health check endpoint."""
    return {"status": "ok", "service": "FILLY"}


# ─── Text Analysis ──────────────────────────────────────────────────


@router.post("/analyze", response_model=AnalyzeResponse)
async def analyze_text(request: AnalyzeRequest):
    """
    Analyze Filipino text for normalization suggestions and grammar errors.

    The text is processed by:
    1. FilNormalizer — dictionary + edit-distance slang/abbreviation detection
    2. GECService — rule-based grammar error detection (stub for GECToR)
    """
    text = request.text.strip()
    if not text:
        return AnalyzeResponse()

    logger.info("Analyzing text (%d chars)", len(text))

    # Run normalization
    normalizer = get_normalizer()
    normalizations = normalizer.normalize(text)

    # Run grammar checking
    gec = get_gec_service()
    grammar_corrections = gec.check(text)

    logger.info(
        "Found %d normalizations, %d grammar corrections",
        len(normalizations),
        len(grammar_corrections),
    )

    return AnalyzeResponse(
        normalizations=normalizations,
        grammar_corrections=grammar_corrections,
    )


# ─── Documents ───────────────────────────────────────────────────────


def _count_words(text: str) -> int:
    """Count words in a text string."""
    stripped = text.strip()
    if not stripped:
        return 0
    return len(stripped.split())


@router.post("/document", response_model=DocumentResponse)
async def create_document(doc: DocumentCreate, db: Session = Depends(get_db)):
    """Create a new document."""
    db_doc = Document(
        title=doc.title,
        content=doc.content,
        word_count=_count_words(doc.content),
    )
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)
    logger.info("Created document id=%d title=%s", db_doc.id, db_doc.title)
    return db_doc


@router.get("/document/{doc_id}", response_model=DocumentResponse)
async def get_document(doc_id: int, db: Session = Depends(get_db)):
    """Retrieve a document by ID."""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc


@router.put("/document/{doc_id}", response_model=DocumentResponse)
async def update_document(
    doc_id: int,
    update: DocumentUpdate,
    db: Session = Depends(get_db),
):
    """Update a document's title and/or content."""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if update.title is not None:
        doc.title = update.title
    if update.content is not None:
        doc.content = update.content
        doc.word_count = _count_words(update.content)

    db.commit()
    db.refresh(doc)
    logger.info("Updated document id=%d", doc.id)
    return doc


@router.get("/documents", response_model=list[DocumentResponse])
async def list_documents(db: Session = Depends(get_db)):
    """List all documents, most recent first."""
    docs = db.query(Document).order_by(Document.updated_at.desc()).all()
    return docs


# ─── File Upload ─────────────────────────────────────────────────────


@router.post("/upload", response_model=UploadResponse)
async def upload_file(file: UploadFile = File(...)):
    """
    Upload a .txt or .docx file and extract its text content.

    Supported formats:
    - .txt (plain text, UTF-8)
    - .docx (Microsoft Word, via python-docx)
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    filename_lower = file.filename.lower()

    if filename_lower.endswith(".txt"):
        content_bytes = await file.read()
        text = content_bytes.decode("utf-8", errors="replace")

    elif filename_lower.endswith(".docx"):
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="python-docx is not installed. Run: pip install python-docx",
            )
        content_bytes = await file.read()
        doc = DocxDocument(BytesIO(content_bytes))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Only .txt and .docx files are accepted.",
        )

    logger.info("Uploaded file '%s' (%d chars extracted)", file.filename, len(text))
    return UploadResponse(text=text)


# ─── Suggestion Actions ─────────────────────────────────────────────


@router.post("/suggestion/{suggestion_id}/accept", response_model=SuggestionActionResponse)
async def accept_suggestion(
    suggestion_id: int,
    db: Session = Depends(get_db),
):
    """Mark a suggestion as accepted."""
    suggestion = db.query(Suggestion).filter(Suggestion.id == suggestion_id).first()
    if not suggestion:
        # Return success even if not found (suggestion may be ephemeral/frontend-only)
        return SuggestionActionResponse(success=True, message="Suggestion accepted")

    suggestion.accepted = True
    suggestion.ignored = False
    db.commit()
    return SuggestionActionResponse(success=True, message="Suggestion accepted")


@router.post("/suggestion/{suggestion_id}/ignore", response_model=SuggestionActionResponse)
async def ignore_suggestion(
    suggestion_id: int,
    db: Session = Depends(get_db),
):
    """Mark a suggestion as ignored."""
    suggestion = db.query(Suggestion).filter(Suggestion.id == suggestion_id).first()
    if not suggestion:
        return SuggestionActionResponse(success=True, message="Suggestion ignored")

    suggestion.ignored = True
    suggestion.accepted = False
    db.commit()
    return SuggestionActionResponse(success=True, message="Suggestion ignored")


# ─── Analytics ───────────────────────────────────────────────────────


@router.get("/analytics", response_model=AnalyticsResponse)
async def get_analytics_global(db: Session = Depends(get_db)):
    """Get global analytics across all documents."""
    return _compute_analytics(db)


@router.get("/analytics/{document_id}", response_model=AnalyticsResponse)
async def get_analytics_for_document(
    document_id: int,
    db: Session = Depends(get_db),
):
    """Get analytics for a specific document."""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return _compute_analytics(db, document_id=document_id)


def _compute_analytics(
    db: Session,
    document_id: int | None = None,
) -> AnalyticsResponse:
    """Compute analytics from stored suggestions."""
    query = db.query(Suggestion)
    if document_id is not None:
        query = query.filter(Suggestion.document_id == document_id)

    suggestions = query.all()

    # Count by type
    norm_suggestions = [s for s in suggestions if s.type == "normalization"]
    gram_suggestions = [s for s in suggestions if s.type == "grammar"]

    norm_accepted = sum(1 for s in norm_suggestions if s.accepted)
    norm_ignored = sum(1 for s in norm_suggestions if s.ignored)
    gram_accepted = sum(1 for s in gram_suggestions if s.accepted)
    gram_ignored = sum(1 for s in gram_suggestions if s.ignored)

    total_issues = len(norm_suggestions) + len(gram_suggestions)
    total_fixed = norm_accepted + gram_accepted
    total_ignored = norm_ignored + gram_ignored

    # Quality score: start at 100, subtract based on unresolved issues
    unresolved = total_issues - total_fixed - total_ignored
    quality_score = max(0, 100 - (unresolved * 5))

    return AnalyticsResponse(
        normalization_stats=NormalizationStats(
            slang_count=len(norm_suggestions),
            abbreviation_count=0,  # Will be categorized in Milestone 3
            spelling_variation_count=0,
            total=len(norm_suggestions),
        ),
        grammar_stats=GrammarStats(
            issues_found=len(gram_suggestions),
            issues_fixed=gram_accepted,
            issues_ignored=gram_ignored,
        ),
        quality_score=quality_score,
    )
